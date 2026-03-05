from __future__ import annotations

import csv
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches


BASE_DIR = Path(__file__).resolve().parents[1]
METRICS_DIR = BASE_DIR / "outputs" / "metrics"
FIGURES_DIR = BASE_DIR / "outputs" / "figures"
REPORTS_DIR = BASE_DIR / "reports"
REPORT_PATH = REPORTS_DIR / "SPH6004_Assignment1_Report.docx"


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def pct(v: str) -> str:
    return f"{float(v) * 100:.2f}%"


def num(v: str, digits: int = 4) -> str:
    return f"{float(v):.{digits}f}"


def add_table_from_dicts(doc: Document, rows: list[dict[str, str]], ordered_cols: list[str]) -> None:
    table = doc.add_table(rows=1, cols=len(ordered_cols))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(ordered_cols):
        hdr[i].text = col

    for row in rows:
        cells = table.add_row().cells
        for i, col in enumerate(ordered_cols):
            cells[i].text = row.get(col, "")


def main() -> None:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    split_rows = read_csv_rows(METRICS_DIR / "split_summary.csv")
    fs_rows = read_csv_rows(METRICS_DIR / "feature_selection_summary.csv")
    metrics_rows = read_csv_rows(METRICS_DIR / "metrics_summary.csv")
    selected_rows = read_csv_rows(METRICS_DIR / "selected_features.csv")

    best_auc_row = max(metrics_rows, key=lambda r: float(r["AUC-ROC"]))
    best_recall_row = max(metrics_rows, key=lambda r: float(r["Recall"]))

    doc = Document()
    doc.add_heading("SPH6004 Assignment 1 Report", level=0)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading("1. Problem Statement", level=1)
    doc.add_paragraph(
        "Objective: predict ICU in-unit mortality (icu_death_flag) using static admission snapshot features from the MIMIC-derived dataset."
    )
    doc.add_paragraph(
        "Clinical note: this is an imbalanced clinical outcome dataset where death prevalence is expected to be lower than survival prevalence."
    )

    doc.add_heading("2. Data and Split", level=1)
    if len(split_rows) == 2:
        train_row = next(r for r in split_rows if r["split"] == "train")
        test_row = next(r for r in split_rows if r["split"] == "test")
        doc.add_paragraph(
            f"Train size: {train_row['rows']}, Test size: {test_row['rows']} (stratified 80/20 split)."
        )
        doc.add_paragraph(
            f"Death rate - Train: {pct(train_row['death_rate'])}, Test: {pct(test_row['death_rate'])}."
        )

    doc.add_heading("3. Leakage-Free Pipeline", level=1)
    doc.add_paragraph(
        "Pipeline update applied based on review: all preprocessors and feature selectors are fitted on training data only, then applied to the test set."
    )
    doc.add_paragraph(
        "Steps: (1) drop identifiers/timestamps and known leakage columns, (2) median/mode imputation, one-hot encoding, scaling, "
        "(3) variance threshold, (4) L1 logistic selection + random forest importance ensemble, (5) model training and hold-out evaluation."
    )

    doc.add_heading("4. Feature Selection Results", level=1)
    add_table_from_dicts(doc, fs_rows, ["stage", "feature_count"])
    top_features = [r["selected_features"] for r in selected_rows[:10]]
    doc.add_paragraph("Top selected features (first 10): " + ", ".join(top_features))

    doc.add_heading("5. Model Performance", level=1)
    display_rows = []
    for r in metrics_rows:
        display_rows.append(
            {
                "Model": r["Model"],
                "Accuracy": num(r["Accuracy"]),
                "Balanced_Accuracy": num(r["Balanced_Accuracy"]),
                "Precision": num(r["Precision"]),
                "Recall": num(r["Recall"]),
                "F1-Score": num(r["F1-Score"]),
                "AUC-ROC": num(r["AUC-ROC"]),
                "PR-AUC": num(r["PR-AUC"]),
            }
        )

    add_table_from_dicts(
        doc,
        display_rows,
        [
            "Model",
            "Accuracy",
            "Balanced_Accuracy",
            "Precision",
            "Recall",
            "F1-Score",
            "AUC-ROC",
            "PR-AUC",
        ],
    )

    doc.add_paragraph(
        f"Best AUC-ROC model: {best_auc_row['Model']} (AUC-ROC={num(best_auc_row['AUC-ROC'])}, PR-AUC={num(best_auc_row['PR-AUC'])})."
    )
    doc.add_paragraph(
        f"Best Recall model: {best_recall_row['Model']} (Recall={num(best_recall_row['Recall'])})."
    )
    doc.add_paragraph(
        "Interpretation: Random Forest gives the strongest ranking performance (AUC/PR-AUC), while Logistic Regression provides the highest sensitivity under current threshold settings."
    )

    doc.add_heading("6. Figures", level=1)
    fi_path = FIGURES_DIR / "feature_importances.png"
    roc_path = FIGURES_DIR / "roc_curves.png"

    if fi_path.exists():
        doc.add_paragraph("Feature importances (Random Forest):")
        doc.add_picture(str(fi_path), width=Inches(6.5))
    if roc_path.exists():
        doc.add_paragraph("ROC curves on hold-out test set:")
        doc.add_picture(str(roc_path), width=Inches(6.5))

    doc.add_heading("7. Conclusion", level=1)
    doc.add_paragraph(
        "The project has been updated to a leakage-free experimental setup with reproducible outputs and cleaner project structure. "
        "Current results are consistent with an imbalanced clinical prediction task and suitable for assignment reporting."
    )

    doc.save(REPORT_PATH)
    print(f"Saved report: {REPORT_PATH}")


if __name__ == "__main__":
    main()
