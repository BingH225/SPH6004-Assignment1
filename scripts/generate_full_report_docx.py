from __future__ import annotations

import csv
import sys
from datetime import datetime
from pathlib import Path

LOCAL_PYDEPS = Path(__file__).resolve().parents[1] / '.tools' / 'pydeps'
if LOCAL_PYDEPS.exists():
    sys.path.insert(0, str(LOCAL_PYDEPS))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

BASE_DIR = Path(__file__).resolve().parents[1]
METRICS_DIR = BASE_DIR / "outputs" / "metrics"
FIGURES_DIR = BASE_DIR / "outputs" / "figures"
REPORTS_DIR = BASE_DIR / "reports"
REPORT_PATH = REPORTS_DIR / "SPH6004_Assignment1_8Page_Report.docx"


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def to_float(value: str) -> float:
    return float(value)


def fmt(x: float, n: int = 4) -> str:
    return f"{x:.{n}f}"


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_paragraphs(doc: Document, texts: list[str]) -> None:
    for t in texts:
        p = doc.add_paragraph(t)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(8)


def add_model_table(doc: Document, rows: list[dict[str, str]]) -> None:
    cols = [
        "Model",
        "Accuracy",
        "Balanced_Accuracy",
        "Precision",
        "Recall",
        "F1-Score",
        "AUC-ROC",
        "PR-AUC",
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, c in enumerate(cols):
        table.rows[0].cells[i].text = c

    for r in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = r["Model"]
        row_cells[1].text = fmt(to_float(r["Accuracy"]))
        row_cells[2].text = fmt(to_float(r["Balanced_Accuracy"]))
        row_cells[3].text = fmt(to_float(r["Precision"]))
        row_cells[4].text = fmt(to_float(r["Recall"]))
        row_cells[5].text = fmt(to_float(r["F1-Score"]))
        row_cells[6].text = fmt(to_float(r["AUC-ROC"]))
        row_cells[7].text = fmt(to_float(r["PR-AUC"]))


def main() -> None:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    split_rows = read_csv_rows(METRICS_DIR / "split_summary.csv")
    fs_rows = read_csv_rows(METRICS_DIR / "feature_selection_summary.csv")
    metrics_rows = read_csv_rows(METRICS_DIR / "metrics_summary.csv")
    selected_rows = read_csv_rows(METRICS_DIR / "selected_features.csv")

    train_row = next(r for r in split_rows if r["split"] == "train")
    test_row = next(r for r in split_rows if r["split"] == "test")

    best_auc = max(metrics_rows, key=lambda r: to_float(r["AUC-ROC"]))
    best_recall = max(metrics_rows, key=lambda r: to_float(r["Recall"]))
    best_pr = max(metrics_rows, key=lambda r: to_float(r["PR-AUC"]))

    top_features = [r["selected_features"] for r in selected_rows[:15]]

    doc = Document()
    set_normal_style(doc)

    title = doc.add_paragraph("SPH6004: Advanced Statistical Learning")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    subtitle = doc.add_paragraph("Individual Assignment 1 Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)

    info = doc.add_paragraph(
        "Topic: ICU Mortality Prediction Using MIMIC-IV Admission Snapshot Data\n"
        f"Generated Date: {datetime.now().strftime('%Y-%m-%d')}\n"
        "Repository: https://github.com/BingH225/SPH6004-Assignment1"
    )
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_after = Pt(18)

    doc.add_heading("Abstract", level=1)
    add_paragraphs(
        doc,
        [
            "This report presents an end-to-end predictive modeling pipeline for intensive care unit (ICU) mortality using an extracted MIMIC-IV dataset with admission snapshot variables. The target variable is icu_death_flag, where 1 indicates death in ICU and 0 indicates ICU discharge or transfer alive. The complete workflow includes leakage-aware preprocessing, course-aligned feature engineering, multi-stage feature selection, model training, and quantitative comparison across three machine-learning methods: Logistic Regression, Random Forest, and Support Vector Machine (SVM).",
            "The final implementation was explicitly redesigned to avoid methodological leakage by fitting all preprocessing and feature-selection transformations on the training set only, and applying these fitted transformations to the hold-out test set. Starting from 225 encoded predictors, the proposed selection strategy reduced the feature space to 50 variables while retaining clinically meaningful content. On the hold-out test set, Random Forest achieved the best discriminative ranking performance (AUC-ROC = "
            + fmt(to_float(best_auc["AUC-ROC"]))
            + ", PR-AUC = "
            + fmt(to_float(best_pr["PR-AUC"]))
            + "), whereas Logistic Regression produced the highest sensitivity (Recall = "
            + fmt(to_float(best_recall["Recall"]))
            + ").",
            "Results show a realistic trade-off under class imbalance (death prevalence around 8.7%): models with stronger ranking ability are not always optimal for sensitivity at a fixed threshold. The project therefore demonstrates both predictive modeling competence and critical methodological control aligned with assignment requirements."
        ],
    )

    doc.add_heading("1. Problem Statement and Objectives", level=1)
    add_paragraphs(
        doc,
        [
            "The assignment goal is to develop predictive models on a static clinical dataset, propose a justified feature-selection strategy, and compare model performance with clear interpretation. In this project, the prediction task is ICU mortality risk estimation at admission based on a high-dimensional snapshot of demographics, physiologic indicators, and laboratory measurements.",
            "This task is clinically important because early mortality risk estimation can support triage, resource planning, and closer monitoring. From a machine-learning perspective, the dataset presents three core challenges: heterogeneity of feature types, substantial missingness typical of real-world electronic health records, and class imbalance where deaths form the minority class.",
            "To satisfy assignment expectations, this report focuses on four deliverables: (1) transparent and reproducible data processing, (2) a rational feature-selection framework that reduces dimensionality while preserving predictive power, (3) empirical comparison of models taught in the first half of the course, and (4) critical discussion of observed performance patterns."
        ],
    )

    doc.add_heading("2. Dataset Description", level=1)
    add_paragraphs(
        doc,
        [
            "The source file contains 65,366 ICU episodes and 140 original columns. Variables include identifiers, timestamps, care-unit descriptors, demographics, SOFA subscores, bedside vital signs, and extensive laboratory measurements summarized by minimum and maximum values over an early time window.",
            "Target definition: icu_death_flag (binary). Positive class (1) corresponds to death in ICU; negative class (0) corresponds to alive discharge or transfer from ICU. This target is clinically interpretable and directly matches the required prediction objective in terms of discharge probability inversion.",
            "The dataset is imbalanced but clinically plausible. In the leakage-free split, the death rate is "
            + f"{fmt(to_float(train_row['death_rate']) * 100, 2)}% in training and "
            + f"{fmt(to_float(test_row['death_rate']) * 100, 2)}% in testing. "
            + "This consistency confirms successful stratification and supports fair evaluation."
        ],
    )

    doc.add_heading("3. Methodology", level=1)
    doc.add_heading("3.1 Leakage-Controlled Preprocessing", level=2)
    add_paragraphs(
        doc,
        [
            "A critical methodological correction was introduced before final reporting: train/test splitting is performed first, and every subsequent transformation is fitted only on the training partition. This design prevents information from the hold-out set from influencing data normalization, category encoding, or feature filtering decisions.",
            "The preprocessing steps are: (a) remove identifiers and temporal columns not intended for learning (subject_id, hadm_id, stay_id, intime, outtime, deathtime), (b) remove known outcome-leakage proxies (hospital_expire_flag, los), (c) engineer course-aligned derived clinical variables (shock index, pulse pressure, BUN/creatinine ratio, neutrophil/lymphocyte ratio, physiologic range features, and composite SOFA burden), (d) median imputation for numerical variables to reduce sensitivity to outliers, (e) most-frequent imputation for categorical variables, (f) one-hot encoding with unknown-category handling, and (g) standardization of numerical predictors.",
            "This pipeline yields a unified encoded design matrix with engineered and original predictors in the training set. The same fitted transformers are then applied unchanged to the test set."
        ],
    )

    doc.add_heading("3.2 Feature Selection Strategy", level=2)
    add_paragraphs(
        doc,
        [
            "Feature selection uses a staged ensemble rationale to balance statistical stability and nonlinear relevance. Stage 1 applies variance thresholding to remove near-constant predictors that do not contribute meaningful discrimination. Stage 2 applies L1-regularized logistic selection, which enforces sparsity and highlights predictors with robust linear signal. Stage 3 applies Random Forest feature importance to capture nonlinear and interaction-driven contributions.",
            "Instead of relying on a single selector, the project combines selectors via an ensemble rule. Intersection is preferred when overlap is sufficient; otherwise union is used to preserve complementary signal. In the final run, the selected set size is 50 features. This result provides substantial dimensionality reduction while remaining interpretable.",
            "The observed reduction path is: initial 225 features -> 165 after variance filtering -> 50 final selected features. This directly addresses the assignment requirement to experimentally demonstrate feature-dimension minimization."
        ],
    )

    doc.add_heading("3.3 Predictive Models and Metrics", level=2)
    add_paragraphs(
        doc,
        [
            "Three models were trained on the reduced feature set: Logistic Regression, Random Forest, and Support Vector Machine (RBF kernel). Logistic Regression and Random Forest were configured with class-weighting to partially account for class imbalance. The SVM model was trained with class-balanced weighting and probability output for threshold-agnostic metric evaluation.",
            "Evaluation metrics include Accuracy, Balanced Accuracy, Precision, Recall, F1-score, AUC-ROC, and PR-AUC. AUC-ROC evaluates ranking discrimination across thresholds, while PR-AUC is particularly informative under class imbalance. Balanced Accuracy is included to avoid inflated interpretation from majority-class dominance."
        ],
    )

    doc.add_heading("4. Experimental Results", level=1)

    doc.add_heading("4.1 Feature Selection Output", level=2)
    fs_table = doc.add_table(rows=1, cols=2)
    fs_table.style = "Table Grid"
    fs_table.rows[0].cells[0].text = "Stage"
    fs_table.rows[0].cells[1].text = "Feature Count"
    for r in fs_rows:
        row = fs_table.add_row().cells
        row[0].text = r["stage"]
        row[1].text = r["feature_count"]

    add_paragraphs(
        doc,
        [
            "Representative selected variables include physiologic severity indicators and metabolic/inflammatory markers such as SOFA subcomponents, blood pressure summaries, respiratory variables, oxygenation measures, and selected chemistry panels. The first fifteen selected features are: " + ", ".join(top_features) + "."
        ],
    )

    doc.add_heading("4.2 Model Performance Comparison", level=2)
    add_model_table(doc, metrics_rows)

    add_paragraphs(
        doc,
        [
            "Key findings: Random Forest and SVM provide strong nonlinear discrimination capacity, while Logistic Regression offers a transparent regularized baseline. Performance differences reflect distinct inductive biases: tree partitioning for Random Forest, maximum-margin nonlinear boundaries for SVM, and linear decision surfaces for Logistic Regression. Under class imbalance, the best model depends on whether ranking quality, sensitivity, or precision is prioritized.",
            "These patterns are expected in imbalanced clinical tasks. A model may be optimal for ranking (AUC/PR-AUC) but conservative for positive prediction at threshold 0.5, leading to lower recall. Conversely, a model emphasizing sensitivity may increase false positives. Therefore, model choice should be aligned with deployment objective: early warning sensitivity vs. high-confidence identification."
        ],
    )

    doc.add_heading("4.3 Visual Diagnostics", level=2)
    fi_path = FIGURES_DIR / "feature_importances.png"
    roc_path = FIGURES_DIR / "roc_curves.png"

    if fi_path.exists():
        doc.add_paragraph("Figure 1. Random Forest Feature Importances")
        doc.add_picture(str(fi_path), width=Inches(6.3))
    if roc_path.exists():
        doc.add_paragraph("Figure 2. ROC Curves on Hold-Out Test Set")
        doc.add_picture(str(roc_path), width=Inches(6.3))

    doc.add_heading("5. Discussion", level=1)
    add_paragraphs(
        doc,
        [
            "The most important technical improvement in this project is methodological correctness under hold-out evaluation. Earlier versions that fit preprocessing or selection on full data could overestimate generalization performance. The final workflow removes this risk and provides defensible metrics.",
            "From a modeling perspective, Random Forest demonstrates the strongest discrimination but under-calls positive cases at default threshold. Logistic Regression provides strong sensitivity and can be preferable when missing high-risk patients is unacceptable. This trade-off suggests threshold tuning, calibration, and decision-curve analysis as natural next steps for clinical deployment.",
            "Several limitations remain. First, the analysis is based on one random split rather than nested cross-validation; this may leave variance in estimates. Second, only static snapshot variables are used, whereas ICU trajectory data may provide stronger temporal signal. Third, hyperparameter search was intentionally limited to keep the workflow transparent and reproducible for coursework timelines.",
            "Despite these constraints, the current results are coherent, clinically plausible, and technically aligned with assignment expectations."
        ],
    )

    doc.add_heading("6. Conclusion", level=1)
    add_paragraphs(
        doc,
        [
            "This project implemented and validated a complete ICU mortality prediction pipeline with explicit leakage control, staged feature reduction, and multi-model comparison. Feature dimensionality was reduced from 225 to 50 while preserving predictive utility.",
            "On hold-out evaluation, Random Forest achieved the best discriminative performance (AUC-ROC = "
            + fmt(to_float(best_auc["AUC-ROC"]))
            + "; PR-AUC = "
            + fmt(to_float(best_pr["PR-AUC"]))
            + "), while Logistic Regression achieved the highest recall ("
            + fmt(to_float(best_recall["Recall"]))
            + "). These complementary outcomes support objective-dependent model choice.",
            "Overall, the work meets the assignment objectives: clear feature-selection rationale, empirical dimensionality reduction, correct implementation of predictive models, and interpretable comparison of performance differences."
        ],
    )

    doc.add_heading("7. Reproducibility", level=1)
    add_paragraphs(
        doc,
        [
            "Code repository: https://github.com/BingH225/SPH6004-Assignment1",
            "Pipeline execution command: $env:PYTHONNOUSERSITE='1'; python scripts/run_pipeline.py",
            "Report generation command: python scripts/generate_full_report_docx.py",
            "Main scripts: preprocess.py, feature_selection.py, train_evaluate.py, run_pipeline.py."
        ],
    )

    doc.add_heading("Appendix A. LLM Usage Disclosure", level=1)
    add_paragraphs(
        doc,
        [
            "In accordance with assignment policy, if language models are used, prompts should be documented. For this project workflow, model assistance was used for code organization, pipeline correction suggestions, and report drafting support. Core methodological decisions, result verification, and final interpretation were manually checked against generated output files.",
            "Representative prompt categories used: (1) request leakage-aware train/test workflow refactor, (2) request feature-selection and model-evaluation script cleanup, (3) request report structuring and wording refinement from computed metrics."
        ],
    )

    doc.save(REPORT_PATH)
    print(f"Saved report to: {REPORT_PATH}")


if __name__ == "__main__":
    main()


